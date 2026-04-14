"""
Merge payment-system and subscription/protection content into the academic proposal.

Reads:  ~/Downloads/QuickConnect - Detailed Proposal (Final).docx
Writes: same path (updated), and copies to docs/QuickConnect - Detailed Proposal (Final).docx

Run from repo root:  python docs/update_detailed_proposal.py
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOWNLOADS = Path.home() / "Downloads" / "QuickConnect - Detailed Proposal (Final).docx"
REPO_DOCS = Path(__file__).resolve().parent / "QuickConnect - Detailed Proposal (Final).docx"


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", *, style: str | None = "Normal"
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def chain_after(
    anchor: Paragraph, segments: list[tuple[str, str]]
) -> Paragraph:
    prev = anchor
    for text, style in segments:
        prev = insert_paragraph_after(prev, text, style=style)
    return prev


def find_paragraph(doc: Document, contains: str) -> Paragraph:
    for p in doc.paragraphs:
        if contains in p.text:
            return p
    raise ValueError(f"No paragraph containing: {contains!r}")


def main() -> None:
    source = DOWNLOADS if DOWNLOADS.is_file() else REPO_DOCS
    if not source.is_file():
        raise SystemExit(
            f"Place 'QuickConnect - Detailed Proposal (Final).docx' in Downloads or {REPO_DOCS.parent}"
        )

    doc = Document(str(source))

    def has_text(sub: str) -> bool:
        return any(sub in p.text for p in doc.paragraphs)

    # --- §4.4.3 Wallet Escrow Payment Flow (after atomicity paragraph) ---
    if has_text("Additional payment behaviour: escrow, disputes, and admin resolution"):
        pass  # already patched
    else:
        anchor_443 = find_paragraph(
            doc,
            "The key design property of this flow is atomicity: because the balance deduction",
        )
        chain_after(
            anchor_443,
            [
                ("", "Normal"),
                (
                    "Additional payment behaviour: escrow, disputes, and admin resolution",
                    "Heading 4",
                ),
                (
                    "Beyond the initial pay-in, each payment row progresses through statuses including held (escrow), released, refunded, failed, and disputed. Real wallet escrow is evidenced by a wallet_transactions row of type payment_hold with direction debit and reference_id equal to the booking id; that link is the authoritative signal that customer funds left the wallet for that booking.",
                    "Normal",
                ),
                (
                    "When both parties confirm satisfaction while funds are held, the auto_release_payment_on_both_confirmed trigger credits the provider and marks the payment released. If a dispute arises while funds are held, administrators resolve it via SECURITY DEFINER RPCs: admin_refund_payment returns escrow to the customer, and admin_release_payment pays the provider from escrow.",
                    "Normal",
                ),
                (
                    "If a customer never paid into escrow but work was completed, a provider may open a dispute that creates a disputed payment record without a prior payment_hold. In that non-escrow scenario, resolving in favour of the provider debits the customer wallet by the claimed amount and credits the provider by the same amount so the ledger remains balanced; resolving in favour of the customer closes the dispute without crediting wallet funds the customer never paid in. The database enforces this distinction by detecting escrow through payment_hold transactions on the booking. Optional admin_dispute_debit_customer supports manual customer debits where platform policy allows negative balances. The Admin Reports interface lists disputed and held payments together with party statements and wallet balances to inform these decisions.",
                    "Normal",
                ),
            ],
        )

    # --- §5.2.4 Wallet and Escrow Payment Module ---
    if has_text("Disputes and RLS: My Bookings merges payment rows"):
        pass
    else:
        anchor_524 = find_paragraph(
            doc,
            "auto_release_payment_on_both_confirmed trigger fires, credits the provider wallet",
        )
        chain_after(
            anchor_524,
            [
                ("", "Normal"),
                (
                    "Disputes and RLS: My Bookings merges payment rows from a direct query because nested selects are sometimes empty under row-level security. Payment disputes notify administrators (including when a disputed row is inserted). Admin Reports invokes admin_refund_payment, admin_release_payment, and optionally admin_dispute_debit_customer; these functions run as SECURITY DEFINER so balances cannot be altered from untrusted clients.",
                    "Normal",
                ),
            ],
        )

    # --- §1.5.2 Out of Scope — elaborate subscription / protection (after subscription bullet) ---
    if has_text("Subscription revenue allocation (planned):"):
        pass
    else:
        anchor_scope = find_paragraph(doc, "Subscription-based provider premium features")
        chain_after(
            anchor_scope,
            [
                (
                    "Subscription revenue allocation (planned): a portion of provider subscription fees may fund a ring-fenced pool to compensate providers when customers default on payment or when full recovery through in-app wallet debits is not achievable, subject to eligibility rules, caps, and audit—complementing, not replacing, ledger-based settlement.",
                    "List Paragraph",
                ),
            ],
        )

    # --- §6.3 Future Work (after Expansion Beyond Botswana) ---
    if has_text("Provider subscriptions and risk protection:"):
        pass
    else:
        anchor_63 = find_paragraph(
            doc,
            "The multi-city architecture and localisation infrastructure provide a replicable template for deployment in other Sub-Saharan African markets.",
        )
        chain_after(
            anchor_63,
            [
                ("", "Normal"),
                ("Provider subscriptions and risk protection:", "Normal"),
                (
                    "A subscription model for providers is planned. Subscription revenue would partially underwrite a protection mechanism: when a customer does not pay or cannot be debited fully, qualified providers could receive compensation from this pool according to transparent rules, while normal escrow and dispute RPCs remain the primary settlement path.",
                    "Normal",
                ),
            ],
        )

    # --- Appendix A: admin dispute bullet (idempotent: only replace legacy one-liner) ---
    for p in doc.paragraphs:
        if p.text.strip().startswith(
            "5. Dispute resolution: trigger admin_refund_payment"
        ):
            p.text = (
                "5. Dispute resolution (Admin > Reports): review disputed and held payments, party statements, and wallets; "
                "use admin_refund_payment or admin_release_payment. For disputes without prior wallet escrow, release debits "
                "the customer and credits the provider by the claimed amount; refund closes in the customer's favour without "
                "crediting funds that were never paid in."
            )
            break

    # --- Appendix A.5 Wallet and Payments: extra steps ---
    if has_text("7. After paying, funds stay in escrow"):
        pass
    else:
        anchor_a5 = find_paragraph(
            doc,
            "6. To withdraw, click Withdraw, enter the amount and mobile money number, and confirm.",
        )
        chain_after(
            anchor_a5,
            [
                (
                    "7. After paying, funds stay in escrow until both parties confirm satisfaction on My Bookings (or an admin resolves a dispute).",
                    "List Paragraph",
                ),
                (
                    "8. If you are a provider and the job is complete but the customer never paid from the wallet, use Report non-payment (dispute) on My Bookings; an administrator will review.",
                    "List Paragraph",
                ),
                (
                    "9. Escrow and confirmations are managed under My Bookings, not only on the Wallet page.",
                    "List Paragraph",
                ),
            ],
        )

    REPO_DOCS.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPO_DOCS))
    print(f"Saved: {REPO_DOCS}")
    try:
        shutil.copyfile(REPO_DOCS, DOWNLOADS)
        print(f"Copied to Downloads: {DOWNLOADS}")
    except OSError as e:
        print(
            f"Note: could not write to Downloads (close Word if open): {e}\n"
            f"Copy manually from: {REPO_DOCS}"
        )


if __name__ == "__main__":
    main()
